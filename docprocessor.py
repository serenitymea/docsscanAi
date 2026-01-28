import json
from pathlib import Path

import docx
import pandas as pd
import PyPDF2
from bs4 import BeautifulSoup


class DocumentProcessor:
    """Universal text extractor for common document formats."""

    SUPPORTED = {
        ".txt", ".pdf", ".docx",
        ".html", ".htm",
        ".csv",
        ".xlsx", ".xls",
        ".json"
    }

    @classmethod
    def extract_text(cls, file_path: str) -> str:
        path = Path(file_path)

        if path.suffix.lower() not in cls.SUPPORTED:
            raise ValueError(f"Unsupported format: {path.suffix}")

        handler = cls._handlers().get(path.suffix.lower())
        if not handler:
            raise ValueError("No handler for this format")

        return handler(path).strip()


    @classmethod
    def _handlers(cls):
        return {
            ".txt": cls._txt,
            ".pdf": cls._pdf,
            ".docx": cls._docx,
            ".html": cls._html,
            ".htm": cls._html,
            ".csv": cls._csv,
            ".xlsx": cls._excel,
            ".xls": cls._excel,
            ".json": cls._json,
        }


    @staticmethod
    def _txt(path: Path) -> str:
        for enc in ("utf-8", "cp1251", "latin1"):
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("Text decoding failed")

    @staticmethod
    def _pdf(path: Path) -> str:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(
                page.extract_text() or ""
                for page in reader.pages
            )

    @staticmethod
    def _docx(path: Path) -> str:
        doc = docx.Document(path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables = [
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for table in doc.tables
            for row in table.rows
        ]

        return "\n".join(paragraphs + tables)

    @staticmethod
    def _html(path: Path) -> str:
        soup = BeautifulSoup(path.read_text(encoding="utf-8"), "html.parser")

        for tag in soup(["script", "style"]):
            tag.decompose()

        return " ".join(soup.stripped_strings)

    @staticmethod
    def _csv(path: Path) -> str:
        df = pd.read_csv(path)
        return df.to_string(index=False)

    @staticmethod
    def _excel(path: Path) -> str:
        sheets = pd.read_excel(path, sheet_name=None)

        return "\n\n".join(
            f"Sheet: {name}\n{df.to_string(index=False)}"
            for name, df in sheets.items()
        )

    @staticmethod
    def _json(path: Path) -> str:
        data = json.loads(path.read_text(encoding="utf-8"))
        return json.dumps(data, indent=2, ensure_ascii=False)
